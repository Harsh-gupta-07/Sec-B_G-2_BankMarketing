from __future__ import annotations

from pathlib import Path
from collections import Counter, defaultdict
import subprocess

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
REPORTS = ROOT / "reports"
ASSETS = REPORTS / "_project_report_assets"
ASSETS.mkdir(parents=True, exist_ok=True)

DATA = pd.read_csv(ROOT / "data" / "processed" / "cleaned_dataset.csv")
SCREENSHOT_DIR = ROOT / "tableau" / "screenshots"

TEAM_MEMBERS = [
    "Harshvardhan Gupta",
    "Kartik Madaan",
    "Dev Tyagi",
    "Shitanshu Tiwari",
    "Satwik Mani Tripathi",
    "Siddharth Dangi",
]

TEAM_ID = "Sec-B G-2"
INSTITUTE = "Newton School of Technology, Rishihood University"
SECTOR = "Finance / Bank Marketing"
PROJECT_TITLE = "Bank Term Deposit Campaign Intelligence"
REPORT_DATE = "29 April 2026"
FONT_REG = "/System/Library/Fonts/Supplemental/Arial.ttf"
FONT_BOLD = "/System/Library/Fonts/Supplemental/Arial Bold.ttf"


def fmt_int(v: float) -> str:
    return f"{int(round(v)):,}"


def fmt_pct(v: float) -> str:
    return f"{v * 100:.2f}%"


def fmt_currency(v: float) -> str:
    return f"€{v:,.0f}"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=100, bottom=80, end=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D9E2EC"):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:color"), color)
        borders.append(el)
    tblPr.append(borders)


def style_document(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor(30, 41, 59)

    for name, size, color in [
        ("Title", 22, "0F172A"),
        ("Heading 1", 16, "0F172A"),
        ("Heading 2", 12, "0E7490"),
        ("Heading 3", 10.5, "1D4ED8"),
    ]:
        st = styles[name]
        st.font.name = "Calibri"
        st.font.bold = True
        st.font.size = Pt(size)
        st.font.color.rgb = RGBColor.from_string(color)


def configure_margins(section):
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)


def add_footer(section, text="Bank Marketing DVA Capstone"):
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)


def add_page_break(doc):
    doc.add_page_break()


def add_paragraph(doc, text, style=None, bold=False, italic=False, align=None, space_after=6):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.15
    return p


def add_bullets(doc, items, level=0, style=None):
    for item in items:
        p = doc.add_paragraph(style=style)
        p.style = doc.styles["List Bullet"]
        if level:
            p.paragraph_format.left_indent = Inches(0.18 * level)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.12
        run = p.add_run(item)
        run.font.name = "Calibri"
        run.font.size = Pt(10.5)


def add_section_heading(doc, title, subtitle=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.name = "Calibri"
    r.font.size = Pt(16)
    r.font.color.rgb = RGBColor(15, 23, 42)
    if subtitle:
        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(8)
        rr = p2.add_run(subtitle)
        rr.font.name = "Calibri"
        rr.font.size = Pt(10.5)
        rr.font.color.rgb = RGBColor(100, 116, 139)


def add_small_label(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = "Calibri"
    r.font.size = Pt(8.5)
    r.bold = True
    r.font.color.rgb = RGBColor(14, 116, 144)


def rgb(hex_color: str):
    hex_color = hex_color.lstrip("#")
    return tuple(int(hex_color[i : i + 2], 16) for i in (0, 2, 4))


def fit_text(draw, text, font, max_width):
    if draw.textlength(text, font=font) <= max_width:
        return text
    words = text.split()
    lines = []
    current = ""
    for word in words:
        trial = word if not current else current + " " + word
        if draw.textlength(trial, font=font) <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)
    return "\n".join(lines[:2])


def make_canvas(width, height, title):
    img = Image.new("RGB", (width, height), "white")
    draw = ImageDraw.Draw(img)
    font_title = ImageFont.truetype(FONT_BOLD, 22)
    draw.text((20, 16), title, fill=rgb("0F172A"), font=font_title)
    return img, draw


def create_bar_chart(path, labels, values, title, color="2F6DAE", horizontal=False, width=680, height=360):
    img, draw = make_canvas(width, height, title)
    font = ImageFont.truetype(FONT_REG, 14)
    font_small = ImageFont.truetype(FONT_REG, 12)
    font_bold = ImageFont.truetype(FONT_BOLD, 13)
    bar_color = rgb(color)
    axis = rgb("334155")
    grid = rgb("E2E8F0")
    plot_left, plot_top, plot_right, plot_bottom = (190, 58, 30, 42) if horizontal else (52, 58, 30, 52)
    plot_w = width - plot_left - plot_right
    plot_h = height - plot_top - plot_bottom
    max_v = max(values) * 1.15 if values else 1
    if horizontal:
        n = len(labels)
        slot = plot_h / n
        for t in range(6):
            frac = t / 5
            x = plot_left + frac * plot_w
            draw.line((x, plot_top, x, plot_top + plot_h), fill=grid, width=1)
            tick = max_v * frac
            tick_label = f"{tick:.0f}%" if max_v <= 100 else f"{tick:,.0f}"
            tw = draw.textlength(tick_label, font=font_small)
            draw.text((x - tw / 2, plot_top + plot_h + 4), tick_label, fill=axis, font=font_small)
        for i, (lab, val) in enumerate(zip(labels, values)):
            y = plot_top + i * slot
            bar_h = slot * 0.64
            cy = y + (slot - bar_h) / 2
            bar_w = (val / max_v) * plot_w
            draw.rounded_rectangle((plot_left, cy, plot_left + bar_w, cy + bar_h), radius=6, fill=bar_color)
            lab_text = fit_text(draw, lab, font, 165)
            draw.multiline_text((20, cy + bar_h / 2 - 10), lab_text, fill=axis, font=font, spacing=2)
            val_label = f"{val:.2f}%" if max_v <= 100 else f"{val:,.0f}"
            draw.text((plot_left + bar_w + 8, cy + bar_h / 2 - 9), val_label, fill=rgb("0F172A"), font=font_bold)
    else:
        n = len(labels)
        slot = plot_w / n
        for t in range(5):
            frac = t / 4
            y = plot_top + plot_h - frac * plot_h
            draw.line((plot_left, y, plot_left + plot_w, y), fill=grid, width=1)
            tick = max_v * frac
            tick_label = f"{tick:.0f}%" if max_v <= 100 else f"{tick:,.0f}"
            tw = draw.textlength(tick_label, font=font_small)
            draw.text((plot_left - tw - 8, y - 8), tick_label, fill=axis, font=font_small)
        for i, (lab, val) in enumerate(zip(labels, values)):
            bar_w = slot * 0.62
            x = plot_left + i * slot + (slot - bar_w) / 2
            bar_h = (val / max_v) * plot_h
            y = plot_top + plot_h - bar_h
            draw.rounded_rectangle((x, y, x + bar_w, plot_top + plot_h), radius=6, fill=bar_color)
            label_text = fit_text(draw, lab, font_small, 60)
            draw.multiline_text((x + bar_w / 2, plot_top + plot_h + 6), label_text, fill=axis, font=font_small, align="center", anchor="ma")
            val_label = f"{val:.2f}%" if max_v <= 100 else f"{val:,.0f}"
            twv = draw.textlength(val_label, font=font_bold)
            draw.text((x + bar_w / 2 - twv / 2, y - 18), val_label, fill=rgb("0F172A"), font=font_bold)
    img.save(path)


def create_line_chart(path, labels, values, title, width=680, height=360):
    img, draw = make_canvas(width, height, title)
    font_small = ImageFont.truetype(FONT_REG, 12)
    font_bold = ImageFont.truetype(FONT_BOLD, 13)
    line_color = rgb("2F6DAE")
    grid = rgb("E2E8F0")
    axis = rgb("334155")
    plot_left, plot_top, plot_right, plot_bottom = 58, 58, 25, 50
    plot_w = width - plot_left - plot_right
    plot_h = height - plot_top - plot_bottom
    max_v = max(values) * 1.15 if values else 1
    for t in range(5):
        frac = t / 4
        y = plot_top + plot_h - frac * plot_h
        draw.line((plot_left, y, plot_left + plot_w, y), fill=grid, width=1)
        tick = max_v * frac
        tick_label = f"{tick:.0f}%" if max_v <= 100 else f"{tick:,.0f}"
        tw = draw.textlength(tick_label, font=font_small)
        draw.text((plot_left - tw - 8, y - 8), tick_label, fill=axis, font=font_small)
    pts = []
    n = len(values)
    for i, val in enumerate(values):
        x = plot_left + (plot_w * i / max(n - 1, 1))
        y = plot_top + plot_h - ((val / max_v) * plot_h)
        pts.append((x, y))
    if len(pts) > 1:
        draw.line(pts, fill=line_color, width=3)
    for x, y in pts:
        draw.ellipse((x - 4, y - 4, x + 4, y + 4), fill=line_color, outline=line_color)
    for i, lab in enumerate(labels):
        x = plot_left + (plot_w * i / max(n - 1, 1))
        label_text = fit_text(draw, lab, font_small, 60)
        draw.multiline_text((x, plot_top + plot_h + 6), label_text, fill=axis, font=font_small, align="center", anchor="ma")
        val = values[i]
        val_label = f"{val:.2f}%"
        twv = draw.textlength(val_label, font=font_bold)
        draw.text((x - twv / 2, plot_top + plot_h - ((val / max_v) * plot_h) - 20), val_label, fill=rgb("0F172A"), font=font_bold)
    img.save(path)


def crosstab_table(df, col):
    rate = (df.groupby(col)["y"].mean() * 100).round(2)
    count = df.groupby(col).size()
    out = pd.DataFrame({"Category": rate.index, "Count": count.values, "Conversion Rate %": rate.values})
    return out.sort_values("Conversion Rate %", ascending=False)


def cramers_v(x, y):
    ct = pd.crosstab(x, y)
    obs = ct.to_numpy(dtype=float)
    n = obs.sum()
    row_sums = obs.sum(axis=1, keepdims=True)
    col_sums = obs.sum(axis=0, keepdims=True)
    expected = row_sums @ col_sums / n
    chi2 = ((obs - expected) ** 2 / expected).sum()
    r, c = obs.shape
    v = (chi2 / (n * (min(r - 1, c - 1)))) ** 0.5
    return chi2, v


def repo_commit_summary():
    out = subprocess.check_output(["git", "log", "--format=%an|%ad|%s", "--date=short"], cwd=ROOT).decode().splitlines()
    by_author = defaultdict(list)
    for line in out:
        author, date, msg = line.split("|", 2)
        by_author[author].append((date, msg))
    return by_author


def build_assets():
    charts = {}
    job = crosstab_table(DATA, "job").head(8)
    create_bar_chart(
        ASSETS / "job_conversion.png",
        job["Category"].tolist(),
        job["Conversion Rate %"].tolist(),
        "Conversion rate by job category",
        color="2F6DAE",
        horizontal=True,
        width=820,
        height=420,
    )
    charts["job"] = ASSETS / "job_conversion.png"

    months = pd.DataFrame({
        "Month": ["jan", "feb", "mar", "apr", "may", "jun", "jul", "aug", "sep", "oct", "nov", "dec"],
        "Rate": (DATA.groupby("month")["y"].mean() * 100).reindex(["jan", "feb", "mar", "apr", "may", "jun", "jul", "aug", "sep", "oct", "nov", "dec"]).round(2).tolist(),
    })
    create_bar_chart(
        ASSETS / "month_conversion.png",
        months["Month"].tolist(),
        months["Rate"].tolist(),
        "Conversion rate by month",
        color="14B8A6",
        horizontal=False,
        width=820,
        height=420,
    )
    charts["month"] = ASSETS / "month_conversion.png"

    campaign = crosstab_table(DATA, "campaign_band")
    create_bar_chart(
        ASSETS / "campaign_conversion.png",
        campaign["Category"].tolist(),
        campaign["Conversion Rate %"].tolist(),
        "Conversion rate by campaign touch band",
        color="E7A33E",
        horizontal=False,
        width=820,
        height=420,
    )
    charts["campaign"] = ASSETS / "campaign_conversion.png"
    return charts


def add_table(doc, df, widths=None, style="Table Grid", shade_header=True, font_size=9, first_row_bold=True):
    table = doc.add_table(rows=len(df) + 1, cols=len(df.columns))
    table.style = style
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    if widths:
        for col_idx, width in enumerate(widths):
            for row in table.rows:
                row.cells[col_idx].width = width
    headers = list(df.columns)
    for j, head in enumerate(headers):
        cell = table.cell(0, j)
        cell.text = str(head)
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(font_size)
                run.font.name = "Calibri"
                run.font.color.rgb = RGBColor(255, 255, 255)
        if shade_header:
            set_cell_shading(cell, "0F172A")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for i, row in enumerate(df.itertuples(index=False), start=1):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            text = str(val)
            cell.text = text
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.0
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 or (widths and widths[j].inches > 1.2 if hasattr(widths[j], "inches") else False) else WD_ALIGN_PARAGRAPH.CENTER
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(font_size)
                    run.font.color.rgb = RGBColor(30, 41, 59)
            if i % 2 == 0:
                set_cell_shading(cell, "F8FAFC")
    set_table_borders(table)
    return table


def add_image(doc, path, width_inches=6.8, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    p.paragraph_format.space_after = Pt(2)
    if caption:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = c.add_run(caption)
        rr.italic = True
        rr.font.size = Pt(8.5)
        rr.font.color.rgb = RGBColor(100, 116, 139)


def build_report():
    charts = build_assets()
    by_author = repo_commit_summary()

    doc = Document()
    style_document(doc)
    section = doc.sections[0]
    configure_margins(section)
    add_footer(section)

    # Cover page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("DATA VISUALIZATION & ANALYTICS")
    r.bold = True
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(14, 116, 144)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(PROJECT_TITLE)
    r.bold = True
    r.font.size = Pt(24)
    r.font.color.rgb = RGBColor(15, 23, 42)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"{SECTOR} | Project Report")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(71, 85, 105)

    add_paragraph(
        doc,
        "This report analyses the UCI Bank Marketing dataset to identify which clients are most likely to subscribe to a term deposit, how campaign design affects conversion, and how Tableau dashboards can support operational targeting decisions.",
        align=WD_ALIGN_PARAGRAPH.CENTER,
        space_after=10,
    )

    facts = pd.DataFrame(
        [
            ["Team ID", TEAM_ID],
            ["Institute", INSTITUTE],
            ["Team members", "; ".join(TEAM_MEMBERS)],
            ["Dataset", "UCI Bank Marketing (bank-full.csv)"],
            ["Rows / columns", f"{fmt_int(len(DATA))} / {len(DATA.columns)} raw columns"],
            ["Submitted", REPORT_DATE],
        ],
        columns=["Field", "Details"],
    )
    add_table(doc, facts, widths=[Inches(1.4), Inches(4.8)], font_size=9.5)

    doc.add_paragraph()
    add_small_label(doc, "At a glance")
    intro = pd.DataFrame(
        [
            ["Total contacts", fmt_int(len(DATA))],
            ["Subscriptions", fmt_int(int(DATA['y'].sum()))],
            ["Subscription rate", fmt_pct(DATA['y'].mean())],
            ["Avg. campaign touches", f"{DATA['campaign'].mean():.2f}"],
        ],
        columns=["Metric", "Value"],
    )
    add_table(doc, intro, widths=[Inches(2.2), Inches(1.4)], font_size=9.5)
    add_paragraph(doc, "The report uses only repository-visible data and dashboard assets stored in this project workspace.", italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)

    # Executive summary
    add_page_break(doc)
    add_section_heading(doc, "1. Executive Summary", "One-page summary for decision makers")
    summary_paras = [
        f"The dataset contains {fmt_int(len(DATA))} client contacts and {fmt_int(int(DATA['y'].sum()))} term-deposit subscriptions, producing an overall subscription rate of {fmt_pct(DATA['y'].mean())}.",
        f"The strongest business signal is prior campaign history: previously contacted clients convert at {fmt_pct(DATA.groupby('previously_contacted')['y'].mean().loc[1])}, compared with {fmt_pct(DATA.groupby('previously_contacted')['y'].mean().loc[0])} for clients without prior contact.",
        f"Segment quality also matters. Students ({DATA.groupby('job')['y'].mean().sort_values(ascending=False).loc['student']*100:.2f}%), retirees ({DATA.groupby('job')['y'].mean().sort_values(ascending=False).loc['retired']*100:.2f}%) and older clients aged 65+ ({DATA.groupby('age_group')['y'].mean().sort_values(ascending=False).loc['65+']*100:.2f}%) respond well.",
        f"Repeated calling is inefficient: conversion falls from {DATA.groupby('campaign_band')['y'].mean().loc['1']*100:.2f}% on first touch to {DATA.groupby('campaign_band')['y'].mean().loc['6+']*100:.2f}% for six or more contacts.",
        f"The report therefore recommends a warm-lead-first call queue, cellular-first outreach, a cap on repeated touches after three attempts, and a stronger focus on high-response months such as March, September, October, and December.",
    ]
    for para in summary_paras:
        add_paragraph(doc, para)
    add_bullets(doc, [
        "Decision rule: prioritize warm leads before cold prospects.",
        "Campaign design rule: keep contact attempts short and targeted.",
        "Channel rule: use cellular-first routing for higher conversion.",
        "Timing rule: schedule the heaviest pushes in the best-performing months.",
    ])

    # Sector context
    add_page_break(doc)
    add_section_heading(doc, "2. Sector Context and Problem Statement", "Finance-sector lead generation is a prioritization problem, not a volume problem.")
    add_paragraph(doc, "Retail banks run large outbound campaigns to sell term deposits. The challenge is that calling everyone wastes capacity, depresses conversion efficiency, and creates noise for low-propensity clients.")
    add_paragraph(doc, "The decision-maker in this project is a campaign or CRM manager who needs a simple rule set: who to call first, how often to call, and which channels should be used for the best return on effort.")
    add_bullets(doc, [
        "Business question: which clients should be contacted first to maximize subscription probability?",
        "Operational question: how can the bank reduce wasted calls while protecting customer experience?",
        "Outcome: a Tableau dashboard and Python analysis that support segmentation-driven outreach.",
    ])

    # Data description
    add_page_break(doc)
    add_section_heading(doc, "3. Data Description", "Source, structure, size, and known limitations")
    add_paragraph(doc, "The analysis uses the UCI Machine Learning Repository's Bank Marketing dataset (`bank-full.csv`). The raw file is semicolon-delimited and has one row per client contact record.")
    data_desc = pd.DataFrame(
        [
            ["Source", "UCI Machine Learning Repository"],
            ["Raw file", "bank-full.csv"],
            ["Rows", fmt_int(len(DATA))],
            ["Raw columns", "17"],
            ["Cleaned columns", str(len(DATA.columns))],
            ["Unit of analysis", "One client contact record"],
            ["Target variable", "y (term-deposit subscription: 0/1)"],
        ],
        columns=["Attribute", "Details"],
    )
    add_table(doc, data_desc, widths=[Inches(2.1), Inches(3.9)], font_size=9.5)
    add_paragraph(doc, "Key limitations: the file uses the label `unknown` in several categorical fields rather than missing values; `duration` is informative for descriptive analysis but leaks information in a predictive setting; and the data represents one historical campaign period, so seasonality should be validated before operational use.")
    add_paragraph(doc, "Important derived fields introduced for Tableau and analysis include `age_group`, `balance_band`, `campaign_band`, `previously_contacted`, and `duration_minutes`.")

    # Cleaning and ETL
    add_page_break(doc)
    add_section_heading(doc, "4. Data Cleaning and ETL Methodology", "Python pipeline from raw file to Tableau-ready dataset")
    add_paragraph(doc, "The ETL pipeline in `scripts/etl_pipeline.py` standardizes the raw file, validates required columns, preserves semantic labels, and adds business-friendly derived columns.")
    add_bullets(doc, [
        "Normalize headers to lowercase snake_case and standardize string values.",
        "Validate expected raw columns before further processing.",
        "Drop duplicates and preserve explicit `unknown` labels rather than silently imputing them.",
        "Create `previously_contacted` from `pdays != -1`.",
        "Bucket `age` into life-stage bands and `balance` into financial bands.",
        "Create `campaign_band` to distinguish first-touch, low-touch, and high-touch campaigns.",
        "Export the cleaned dataset and Tableau-ready subset with 22 columns.",
    ])
    etl_steps = pd.DataFrame(
        [
            ["Input", "data/raw/bank-full.csv", "Raw UCI source"],
            ["Standardize", "Lowercase / snake_case / strip spaces", "Consistent joins and filtering"],
            ["Derive", "age_group, balance_band, campaign_band, etc.", "Business segmentation"],
            ["Validate", "Required columns and yes/no target values", "Data quality control"],
            ["Output", "cleaned_dataset.csv / Tableau-ready CSV", "Analysis and dashboard consumption"],
        ],
        columns=["Stage", "What happens", "Why it matters"],
    )
    add_table(doc, etl_steps, widths=[Inches(1.0), Inches(3.0), Inches(1.8)], font_size=8.6)

    # KPI framework
    add_page_break(doc)
    add_section_heading(doc, "5. KPI and Metric Framework", "Metrics used to judge conversion, targeting quality, and financial context")
    kpis = pd.DataFrame(
        [
            ["Subscription rate", fmt_pct(DATA['y'].mean()), "Subscribers / total contacts", "North-star conversion metric"],
            ["Total subscriptions", fmt_int(int(DATA['y'].sum())), "Count of y = 1 records", "Absolute campaign success"],
            ["Avg. contacts / client", f"{DATA['campaign'].mean():.2f}", "Mean campaign touch count", "Effort intensity"],
            ["Previously contacted share", fmt_pct(DATA['previously_contacted'].mean()), "previously_contacted = 1", "Warm-lead prevalence"],
            ["Avg. yearly balance", fmt_currency(DATA['balance'].mean()), "Mean balance across all clients", "Financial capability"],
            ["Default rate", fmt_pct((DATA['default'] == 'yes').mean()), "Clients with default = yes", "Credit risk context"],
            ["Loan rate", fmt_pct((DATA['loan'] == 'yes').mean()), "Clients with loan = yes", "Debt burden context"],
        ],
        columns=["KPI", "Current value", "Definition", "Business meaning"],
    )
    add_table(doc, kpis, widths=[Inches(1.5), Inches(1.1), Inches(2.0), Inches(1.5)], font_size=8.4)
    add_paragraph(doc, "For campaign management, the most useful pairings are subscription rate versus contact intensity, and subscription rate versus prior contact history. These pairs reveal where the bank is gaining conversion and where it is wasting touches.")

    # EDA page
    add_page_break(doc)
    add_section_heading(doc, "6. Exploratory Data Analysis", "Visual patterns that shape segmentation and targeting")
    add_paragraph(doc, "The EDA shows that subscription likelihood is not uniform. Certain customer groups respond materially better, and some outreach patterns are clearly inefficient.")
    add_image(doc, charts["job"], width_inches=6.7, caption="Figure 1. Conversion rate by job category.")
    add_paragraph(doc, "Students and retirees lead the conversion table, while blue-collar, entrepreneur, and housemaid segments sit near the bottom. That makes occupation a practical prioritization filter rather than a descriptive label.")
    add_image(doc, charts["month"], width_inches=6.7, caption="Figure 2. Seasonality of conversion across months.")
    add_paragraph(doc, "Campaign response spikes in March, December, September, and October. Timing matters, so calendar windows should be treated as a campaign lever rather than an afterthought.")

    # Statistical analysis
    add_page_break(doc)
    add_section_heading(doc, "7. Statistical Analysis Results", "Association strength and decision relevance")
    add_paragraph(doc, "The goal of this section is not to produce a machine-learning model, but to quantify which variables are most strongly associated with the subscription outcome.")
    stats_rows = []
    for col in ["poutcome", "previously_contacted", "contact", "age_group", "job", "balance_band", "campaign_band", "education", "marital"]:
        chi2, v = cramers_v(DATA[col], DATA["y"])
        stats_rows.append([col, f"{chi2:,.1f}", f"{v:.3f}", ""])
    stats_df = pd.DataFrame(stats_rows, columns=["Driver", "Chi-square", "Cramér's V", "Interpretation"])
    stats_df.loc[stats_df["Driver"] == "poutcome", "Interpretation"] = "Strongest relationship; prior success is highly informative."
    stats_df.loc[stats_df["Driver"] == "previously_contacted", "Interpretation"] = "Warm leads materially outperform cold leads."
    stats_df.loc[stats_df["Driver"] == "contact", "Interpretation"] = "Cellular contact clearly outperforms telephone."
    stats_df.loc[stats_df["Driver"] == "age_group", "Interpretation"] = "Older bands convert materially better."
    stats_df.loc[stats_df["Driver"] == "job", "Interpretation"] = "Occupation is a useful segmentation feature."
    stats_df.loc[stats_df["Driver"] == "balance_band", "Interpretation"] = "Higher balances align with better conversion."
    stats_df.loc[stats_df["Driver"] == "campaign_band", "Interpretation"] = "Conversion decays as touches increase."
    stats_df.loc[stats_df["Driver"] == "education", "Interpretation"] = "Education matters, but less than contact history."
    stats_df.loc[stats_df["Driver"] == "marital", "Interpretation"] = "Useful as a secondary segmentation feature."
    add_table(doc, stats_df, widths=[Inches(1.7), Inches(1.0), Inches(1.0), Inches(2.7)], font_size=8.2)
    add_paragraph(doc, "Key descriptive contrasts: subscribers have a higher average balance ({0}) than non-subscribers ({1}), and they receive fewer campaign touches on average ({2:.2f} versus {3:.2f}).".format(
        fmt_currency(DATA.loc[DATA['y'] == 1, 'balance'].mean()),
        fmt_currency(DATA.loc[DATA['y'] == 0, 'balance'].mean()),
        DATA.loc[DATA['y'] == 1, 'campaign'].mean(),
        DATA.loc[DATA['y'] == 0, 'campaign'].mean(),
    ))
    add_paragraph(doc, f"Previously contacted clients convert at {DATA.groupby('previously_contacted')['y'].mean().loc[1]*100:.2f}% versus {DATA.groupby('previously_contacted')['y'].mean().loc[0]*100:.2f}% for clients with no prior contact, a lift of roughly {((DATA.groupby('previously_contacted')['y'].mean().loc[1] / DATA.groupby('previously_contacted')['y'].mean().loc[0]) - 1) * 100:.1f}%.")

    # Dashboard design
    add_page_break(doc)
    add_section_heading(doc, "8. Dashboard Design", "Screenshots and what each view is meant to answer")
    add_paragraph(doc, "The workbook is intentionally split into an executive view and operational views so that leadership can read one story while campaign owners can drill into targeting rules.")
    add_image(doc, SCREENSHOT_DIR / "executive_view.png", width_inches=6.8, caption="Dashboard 1. Executive overview.")
    add_bullets(doc, [
        "Shows headline KPIs, high-level segments, and balance distribution.",
        "Best for leadership reporting and fast portfolio review.",
        "Main filters: job, age group, education, marital status, balance band.",
    ])

    add_page_break(doc)
    add_section_heading(doc, "8. Dashboard Design (continued)", "Operational lens one")
    add_image(doc, SCREENSHOT_DIR / "campaign strategy view.png", width_inches=6.8, caption="Dashboard 2. Contact strategy and conversion lens.")
    add_bullets(doc, [
        "Focuses on contact volume, channel performance, prior-contact lift, and campaign touch effects.",
        "Best for CRM managers deciding how to route the next batch of leads.",
        "Main filters: number of contacts, prior contact, channel, campaign history.",
    ])

    add_page_break(doc)
    add_section_heading(doc, "8. Dashboard Design (continued)", "Operational lens two")
    add_image(doc, SCREENSHOT_DIR / "risk_balance_view.png", width_inches=6.8, caption="Dashboard 3. Risk and balance lens.")
    add_bullets(doc, [
        "Brings credit risk context into the story through default, loan, and balance bands.",
        "Best for balancing conversion uplift against customer quality and risk exposure.",
        "Main filters: age group, balance band, loan, default, and outcome.",
    ])

    # Insights and recommendations
    add_page_break(doc)
    add_section_heading(doc, "9. Key Insights and Recommendations", "Decision language, not chart captions")
    insights = pd.DataFrame(
        [
            ["1", "Prioritize students and retirees first.", "They deliver the highest conversion and deserve the first call queue."],
            ["2", "Treat 65+ and 18-24 as high-response age bands.", "Age is a practical targeting dimension, not just a descriptive one."],
            ["3", "Use prior contact history as a primary routing rule.", "Warm leads convert far better than cold leads."],
            ["4", "Cap repeated outreach after three touches.", "Conversion falls sharply after touch count increases."],
            ["5", "Prefer cellular-first contact plans.", "Cellular outperforms telephone and unknown channels are weak."],
            ["6", "Push campaigns in March, September, October, and December.", "Response is materially stronger in these windows."],
            ["7", "Give medium and high balance clients higher priority.", "Higher balance bands outperform low and negative bands."],
            ["8", "Use prior success as a premium routing signal.", "Successful previous outcomes dominate all other signals."],
            ["9", "De-emphasize blue-collar and services segments for premium outreach.", "These groups sit near the bottom of the response table."],
            ["10", "Keep unknown contact records out of prime call queues.", "Missing channel context lowers conversion and should be cleaned."],
        ],
        columns=["#", "Insight", "Business implication"],
    )
    add_table(doc, insights, widths=[Inches(0.4), Inches(2.9), Inches(2.7)], font_size=8.2)
    add_paragraph(doc, "Recommended actions:")
    recs = pd.DataFrame(
        [
            ["Prioritize warm leads", "Route previously contacted leads and prior-success records first.", "Highest immediate conversion gain"],
            ["Apply touch caps", "Stop repeating low-yield calls after three touches.", "Lower wasted call volume"],
            ["Use cellular-first routing", "Make cellular the default channel where possible.", "Stronger response rate"],
            ["Seasonal campaign bursts", "Concentrate spend in peak months.", "Better calendar efficiency"],
        ],
        columns=["Recommendation", "What to do", "Expected impact"],
    )
    add_table(doc, recs, widths=[Inches(1.35), Inches(3.15), Inches(1.35)], font_size=8.2)

    # Impact, limitations, future scope
    add_page_break(doc)
    add_section_heading(doc, "10. Impact Estimation, Limitations, and Future Scope", "What this work can realistically change")
    add_paragraph(doc, f"If 10,000 contacts are routed from the cold baseline to previously contacted leads, the expected incremental lift is about {int(round(10_000 * (DATA.groupby('previously_contacted')['y'].mean().loc[1] - DATA.groupby('previously_contacted')['y'].mean().loc[0]))):,} extra subscriptions.")
    add_paragraph(doc, f"Using a conservative illustrative contribution of INR 50 per incremental subscription, the same shift implies roughly INR {round(10_000 * (DATA.groupby('previously_contacted')['y'].mean().loc[1] - DATA.groupby('previously_contacted')['y'].mean().loc[0]) * 50):,.0f} in value per 10,000 contacted leads.")
    add_bullets(doc, [
        "Limitation 1: no call-cost or offer-cost data, so ROI is estimated rather than exact.",
        "Limitation 2: `duration` is useful for analysis but leaks outcome information and should not be used as a pre-call predictor.",
        "Limitation 3: the dataset captures one campaign period, so seasonality should be revalidated before deployment.",
        "Limitation 4: unknown categories are explicit labels, not absent rows, which is useful for analysis but imperfect for modeling.",
    ])
    add_paragraph(doc, "Future scope should include a cost-per-touch model, a propensity scoring layer, and A/B-tested routing rules that compare warm-lead targeting against broader outreach mixes.")

    # Contribution matrix
    add_page_break(doc)
    add_section_heading(doc, "11. Contribution Matrix")
    commit_counts = Counter({author: len(commits) for author, commits in by_author.items()})
    contrib = pd.DataFrame(
        [
            ["Harshvardhan Gupta", f"{commit_counts.get('Harshvardhan Gupta', 0)} commits", "Raw data, cleaning, EDA, statistical analysis, workbook updates"],
            ["Kartik Madaan", f"{commit_counts.get('madaankartik', 0)} commits", "Statistical analysis, Tableau dashboard, report writing, PPT"],
            ["Dev Tyagi", "4 commits", "Tableau dashboard"],
            ["Shitanshu Tiwari", "2 commits", "EDA and analysis"],
            ["Satwik Mani Tripathi", "No commit", "N/A"],
            ["Siddharth Dangi", "No commit", "N/A"],
        ],
        columns=["Team member", "Git history status", "Repository-visible contribution"],
    )
    add_table(doc, contrib, widths=[Inches(1.55), Inches(1.35), Inches(3.0)], font_size=7.4)

    # Final QA notes page? maybe not needed, keep within 11-12 pages.

    out = REPORTS / "project_report.docx"
    doc.save(out)
    return out


if __name__ == "__main__":
    path = build_report()
    print(path)
